"""
Resume Service - Resume Parsing and ATS Analysis
FastAPI-based microservice for resume management and analysis
"""

import os
import sys
import json
import asyncio
import re
from pathlib import Path
from datetime import datetime
from typing import Optional, List, Dict, Any
from uuid import UUID, uuid4
from io import BytesIO

# Ensure project root is importable when running this file directly.
PROJECT_ROOT = Path(__file__).resolve().parents[2]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from fastapi import FastAPI, HTTPException, Depends, BackgroundTasks, UploadFile, File, Form, status
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from pydantic import BaseModel, Field
from sqlalchemy.ext.asyncio import AsyncSession, create_async_engine, async_sessionmaker
from sqlalchemy import select, update, desc
from sqlalchemy.sql import func
import redis.asyncio as redis

# Import PDF and DOCX parsing libraries
try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    import docx
except ImportError:
    docx = None

try:
    from openai import AsyncOpenAI
except ImportError:
    AsyncOpenAI = None

from database.models import (
    Resume, ResumeSection, ResumeSkill, Skill, User,
    ResumeSourceType, ProficiencyLevel, ResumeAnalysis
)

# ==================== Configuration ====================

DATABASE_URL = os.getenv("DATABASE_URL", "postgresql+asyncpg://user:pass@localhost:5432/aimock")
REDIS_URL = os.getenv("REDIS_URL", "redis://localhost:6379/0")

# ── API key resolution (env var wins; fall back to local api.txt for dev) ──
def _load_api_key() -> str:
    # 1. Environment variable (Docker / production / CI)
    key = os.getenv("OPENROUTER_API_KEY", "").strip()
    if key:
        return key

    # 2. Local api.txt relative to this file (developer convenience)
    local_txt = Path(__file__).parent / "api.txt"
    if local_txt.exists():
        key = local_txt.read_text(encoding="utf-8").strip()
        if key:
            return key

    return ""

OPENROUTER_API_KEY = _load_api_key()
OPENROUTER_BASE_URL = "https://openrouter.ai/api/v1"
# Model name must use OpenRouter's namespaced format
AI_MODEL = "openai/gpt-4o-mini"

FRONTEND_ORIGINS = os.getenv(
    "FRONTEND_ORIGINS",
    "http://localhost:3000,http://127.0.0.1:3000,http://localhost:3001,http://127.0.0.1:3001"
)

# Allowed file extensions for resume uploads
ALLOWED_EXTENSIONS = {'.pdf', '.doc', '.docx'}

# Validation constants
MIN_TEXT_LENGTH = 200  # Minimum extracted text characters
MIN_WORD_COUNT = 50    # Minimum words in resume
MIN_SECTION_WORDS = 30 # Minimum words per section

# ==================== Database Setup ====================

engine = create_async_engine(DATABASE_URL, echo=False, pool_pre_ping=True)
async_session = async_sessionmaker(engine, class_=AsyncSession, expire_on_commit=False)

async def get_db() -> AsyncSession:
    async with async_session() as session:
        try:
            yield session
            await session.commit()
        except Exception:
            await session.rollback()
            raise

# ==================== Redis Setup ====================

redis_client: Optional[redis.Redis] = None

async def get_redis() -> redis.Redis:
    global redis_client
    if redis_client is None:
        redis_client = redis.from_url(REDIS_URL, decode_responses=True)
    return redis_client

# ==================== Pydantic Models ====================

class ResumeCreate(BaseModel):
    source_type: ResumeSourceType
    source_url: Optional[str] = None

class ResumeSectionUpdate(BaseModel):
    section_type: str
    content: Dict[str, Any]
    display_order: Optional[int] = None

class ResumeAnalyzeRequest(BaseModel):
    resume_id: UUID
    target_role: Optional[str] = None

class ResumeParseRequest(BaseModel):
    raw_content: str

class ResumeAnalyzePayload(BaseModel):
    target_role: Optional[str] = None

class ATSAnalysis(BaseModel):
    overall_score: float
    sections: Dict[str, float]
    keyword_match: float
    formatting_score: float
    experience_relevance: float
    education_match: float
    improvements: List[Dict[str, Any]]
    competitive_percentile: float

class ResumeResponse(BaseModel):
    id: UUID
    user_id: UUID
    source_type: ResumeSourceType
    parsed_data: Optional[Dict[str, Any]]
    ats_score: Optional[float]
    version: int
    is_active: bool
    created_at: datetime

class ResumeHistoryResponse(BaseModel):
    id: str
    file_name: str
    ats_score: float
    skills_detected: List[str]
    missing_skills: List[str]
    strengths: List[str]
    improvements: List[str]
    suggestions: List[str]
    uploaded_at: datetime

# ==================== Resume Service ====================

class ResumeService:
    def __init__(self):
        self.openai_key = OPENROUTER_API_KEY

    def _get_ai_client(self) -> Optional[Any]:
        """Return an OpenRouter-compatible AsyncOpenAI client, or None."""
        if not self.openai_key or AsyncOpenAI is None:
            return None
        return AsyncOpenAI(
            api_key=self.openai_key,
            base_url=OPENROUTER_BASE_URL,
        )

    def validate_file_type(self, filename: str) -> bool:
        """Validate that the file is a valid resume format (PDF, DOC, DOCX)"""
        if not filename:
            return False
        ext = Path(filename).suffix.lower()
        return ext in ALLOWED_EXTENSIONS

    def extract_text_from_file(self, content: bytes, filename: str) -> str:
        """Extract text from supported file types. Returns empty string on failure."""
        ext = Path(filename).suffix.lower()

        if ext == '.pdf':
            return self._extract_from_pdf(content)
        elif ext in ['.docx', '.doc']:
            return self._extract_from_docx(content)
        else:
            return ""

    def _extract_from_pdf(self, content: bytes) -> str:
        """Extract text from PDF using pypdf."""
        if PdfReader is None:
            return ""
        try:
            pdf_file = BytesIO(content)
            pdf_reader = PdfReader(pdf_file)
            text_parts: List[str] = []
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            return "\n".join(text_parts).strip()
        except Exception as e:
            print(f"PDF extraction error: {e}")
            return ""

    def _extract_from_docx(self, content: bytes) -> str:
        """Extract text from DOCX using python-docx"""
        if docx is None:
            return ""
        try:
            doc_file = BytesIO(content)
            document = docx.Document(doc_file)
            text_parts = []
            for para in document.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            return "\n".join(text_parts).strip()
        except Exception as e:
            print(f"DOCX extraction error: {e}")
            return ""

    async def parse_resume(self, resume_id: UUID, raw_content: str) -> Dict[str, Any]:
        """Parse resume content via OpenRouter AI. No mock fallback."""
        if not raw_content or len(raw_content.strip()) < MIN_TEXT_LENGTH:
            raise HTTPException(
                status_code=400,
                detail="Could not extract readable text from file. Ensure it's a text-based PDF/DOCX, not a scanned image or encrypted file."
            )

        client = self._get_ai_client()
        if client is None:
            raise HTTPException(
                status_code=503,
                detail="Resume parsing is unavailable: OPENROUTER_API_KEY is not configured."
            )

        truncated_content = raw_content[:8000]
        prompt = (
            "Extract resume data from the text below. Return ONLY a valid JSON object with keys: "
            "contact_info, professional_summary, skills, experience, education, projects. "
            "Do not invent facts. Use empty string/empty array when data is missing.\n\n"
            f"Resume text:\n{truncated_content}"
        )

        try:
            completion = await client.chat.completions.create(
                model=AI_MODEL,
                messages=[
                    {"role": "system", "content": "You are a resume parser that outputs strict JSON only."},
                    {"role": "user", "content": prompt},
                ],
                response_format={"type": "json_object"},
                temperature=0,
            )
            raw_json = completion.choices[0].message.content or "{}"
            parsed_data = json.loads(raw_json)
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(status_code=502, detail=f"Resume parsing failed: {str(e)}")

        if not isinstance(parsed_data, dict):
            raise HTTPException(status_code=502, detail="Resume parser returned invalid format.")

        # Normalise fields
        if not isinstance(parsed_data.get("skills", []), list):
            parsed_data["skills"] = []
        if not isinstance(parsed_data.get("experience", []), list):
            parsed_data["experience"] = []
        if not isinstance(parsed_data.get("education", []), list):
            parsed_data["education"] = []
        if not isinstance(parsed_data.get("projects", []), list):
            parsed_data["projects"] = []
        if not isinstance(parsed_data.get("contact_info", {}), dict):
            parsed_data["contact_info"] = {}
        if not isinstance(parsed_data.get("professional_summary", ""), str):
            parsed_data["professional_summary"] = ""

        return parsed_data

    def _extract_skill_names(self, parsed_data: Dict[str, Any]) -> List[str]:
        """Normalize skills from LLM output into a deduplicated list of skill names."""
        skills = parsed_data.get("skills", [])
        names: List[str] = []
        if not isinstance(skills, list):
            return names

        for skill in skills:
            if isinstance(skill, dict):
                name = str(skill.get("name", "")).strip()
            else:
                name = str(skill).strip()
            if name and name.lower() not in [s.lower() for s in names]:
                names.append(name)
        return names

    async def generate_professional_ats_analysis(
        self,
        parsed_data: Dict[str, Any],
        target_role: Optional[str] = None
    ) -> Dict[str, Any]:
        """Generate a professional structured ATS analysis"""

        target = target_role or "software engineer"
        required_keywords = self._get_role_keywords(target)

        skills_detected = self._extract_skill_names(parsed_data)

        missing_skills = [kw for kw in required_keywords
                          if kw.lower() not in [s.lower() for s in skills_detected]]

        strengths = []
        if parsed_data.get("experience") and len(parsed_data.get("experience", [])) > 0:
            strengths.append("Work experience section present")
        if parsed_data.get("education") and len(parsed_data.get("education", [])) > 0:
            strengths.append("Education section present")
        if skills_detected:
            strengths.append(f"Listed {len(skills_detected)} technical skills")
        if parsed_data.get("projects") and len(parsed_data.get("projects", [])) > 0:
            strengths.append("Projects section showcases practical work")
        if parsed_data.get("professional_summary") and parsed_data.get("professional_summary") != "Not found":
            strengths.append("Professional summary provides quick overview")

        improvements = []
        if len(skills_detected) < 5:
            improvements.append("Add more technical skills relevant to your target role")
        if not parsed_data.get("projects"):
            improvements.append("Add a projects section to showcase practical work")
        if not parsed_data.get("professional_summary") or parsed_data.get("professional_summary") == "Not found":
            improvements.append("Add a professional summary at the top of your resume")
        if missing_skills:
            improvements.append(f"Consider adding: {', '.join(missing_skills[:3])}")

        suggestions = []
        if missing_skills:
            suggestions.append(f"Include missing keywords: {', '.join(missing_skills[:5])}")
        suggestions.append("Use standard section headings (Experience, Education, Skills)")
        suggestions.append("Keep formatting consistent throughout the resume")
        if parsed_data.get("experience"):
            suggestions.append("Quantify achievements with numbers and percentages")
        if skills_detected:
            suggestions.append("Organize skills by category (Languages, Frameworks, Tools)")

        ats_result = await self.calculate_ats_score(parsed_data, target_role)

        return {
            "atsScore": ats_result.get("overall", 0),
            "skillsDetected": skills_detected,
            "missingSkills": missing_skills[:10],
            "strengths": strengths,
            "improvements": improvements,
            "suggestions": suggestions,
            "professionalFeedback": _generate_analysis_summary(ats_result.get("overall", 0)),
            "detailedAnalysis": ats_result
        }

    async def calculate_ats_score(
        self,
        parsed_data: Dict[str, Any],
        target_role: Optional[str] = None
    ) -> Dict[str, Any]:
        """Calculate ATS score with detailed breakdown"""

        scores = {}
        improvements = []

        keyword_score, keyword_improvements = self._score_keywords(parsed_data, target_role)
        scores['keyword_match'] = keyword_score * 0.30
        improvements.extend(keyword_improvements)

        format_score, format_improvements = self._score_formatting(parsed_data)
        scores['formatting'] = format_score * 0.20
        improvements.extend(format_improvements)

        exp_score, exp_improvements = self._score_experience(parsed_data)
        scores['experience'] = exp_score * 0.25
        improvements.extend(exp_improvements)

        edu_score, edu_improvements = self._score_education(parsed_data, target_role)
        scores['education'] = edu_score * 0.10
        improvements.extend(edu_improvements)

        complete_score, complete_improvements = self._score_completeness(parsed_data)
        scores['completeness'] = complete_score * 0.10
        improvements.extend(complete_improvements)

        impact_score, impact_improvements = self._score_impact(parsed_data)
        scores['impact'] = impact_score * 0.05
        improvements.extend(impact_improvements)

        overall = sum(scores.values())

        if overall >= 0.8:
            percentile = 85.0
        elif overall >= 0.6:
            percentile = 65.0
        elif overall >= 0.4:
            percentile = 45.0
        else:
            percentile = 25.0

        return {
            "overall": round(overall * 100, 1),
            "sections": {k: round(v * 100, 1) for k, v in scores.items()},
            "keyword_match": round(keyword_score * 100, 1),
            "formatting_score": round(format_score * 100, 1),
            "experience_relevance": round(exp_score * 100, 1),
            "education_match": round(edu_score * 100, 1),
            "improvements": self._prioritize_improvements(improvements, overall),
            "competitive_percentile": percentile
        }

    def _score_keywords(self, parsed_data, target_role) -> tuple:
        required_keywords = self._get_role_keywords(target_role or "software engineer")
        found_keywords = set()
        resume_text = json.dumps(parsed_data).lower()
        for keyword in required_keywords:
            if keyword.lower() in resume_text:
                found_keywords.add(keyword)
        match_ratio = len(found_keywords) / len(required_keywords) if required_keywords else 1.0
        score = min(match_ratio, 1.0)
        feedback = []
        if match_ratio < 0.5:
            missing = set(required_keywords) - found_keywords
            feedback.append({
                "issue": f"Add missing key skills: {', '.join(list(missing)[:5])}",
                "category": "keyword", "impact": 8, "priority": "high"
            })
        return score, feedback

    def _score_formatting(self, parsed_data: Dict) -> tuple:
        score = 0.85
        feedback = []
        required_sections = ["contact_info", "experience", "education"]
        missing = [s for s in required_sections if s not in parsed_data]
        if missing:
            score -= 0.1 * len(missing)
            feedback.append({
                "issue": f"Missing sections: {', '.join(missing)}",
                "category": "format", "impact": 6, "priority": "medium"
            })
        return max(0, score), feedback

    def _score_experience(self, parsed_data: Dict) -> tuple:
        score = 0.8
        feedback = []
        experiences = parsed_data.get("experience", [])
        if not experiences:
            return 0.0, [{"issue": "Add work experience", "category": "content", "impact": 10, "priority": "critical"}]
        for exp in experiences:
            has_metrics = any(
                c.isdigit() or "$" in c or "%" in c
                for c in exp.get("key_achievements", [])
            )
            if not has_metrics:
                score -= 0.1
                feedback.append({
                    "issue": f"Add quantified achievements at {exp.get('company')}",
                    "category": "content", "impact": 5, "priority": "medium"
                })
        return max(0, min(score, 1)), feedback

    def _score_education(self, parsed_data: Dict, target_role: Optional[str]) -> tuple:
        education = parsed_data.get("education", [])
        if not education:
            return 0.5, [{"issue": "Add education section", "category": "content", "impact": 5, "priority": "medium"}]
        return 0.9, []

    def _score_completeness(self, parsed_data: Dict) -> tuple:
        required = ["contact_info", "professional_summary", "skills", "experience", "education"]
        present = [r for r in required if r in parsed_data and parsed_data[r]]
        score = len(present) / len(required)
        feedback = []
        if score < 1.0:
            missing = set(required) - set(present)
            feedback.append({
                "issue": f"Complete missing sections: {', '.join(missing)}",
                "category": "completeness", "impact": 7, "priority": "high"
            })
        return score, feedback

    def _score_impact(self, parsed_data: Dict) -> tuple:
        score = 0.7
        feedback = []
        experiences = parsed_data.get("experience", [])
        total_achievements = sum(len(e.get("key_achievements", [])) for e in experiences)
        if total_achievements < 3:
            score -= 0.2
            feedback.append({
                "issue": "Add more quantified achievements (%, $, time saved)",
                "category": "enhance", "impact": 4, "priority": "low"
            })
        return max(0, score), feedback

    def _get_role_keywords(self, role: str) -> List[str]:
        common_keywords = {
            "software engineer": [
                "Python", "JavaScript", "Java", "React", "Node.js", "AWS", "SQL",
                "Git", "REST API", "Agile", "Microservices", "Docker", "Kubernetes"
            ],
            "data scientist": [
                "Python", "Machine Learning", "TensorFlow", "PyTorch", "SQL",
                "Statistics", "Data Analysis", "Pandas", "NumPy", "Visualization"
            ],
            "devops engineer": [
                "AWS", "Azure", "Docker", "Kubernetes", "Jenkins", "Terraform",
                "CI/CD", "Linux", "Scripting", "Monitoring", "Ansible"
            ]
        }
        return common_keywords.get(role.lower(), common_keywords["software engineer"])

    def _prioritize_improvements(self, feedback: List[Dict], overall_score: float) -> List[Dict]:
        impact_weights = {
            "missing": 10, "keyword": 8, "format": 6, "content": 5, "enhance": 3
        }
        for item in feedback:
            category = item.get("category", "content")
            item["impact_score"] = impact_weights.get(category, 5)
            item["effort_estimate"] = "medium"
            item["roi_score"] = item["impact_score"] / 2
        feedback.sort(key=lambda x: x.get("roi_score", 0), reverse=True)
        return feedback[:10]


# ==================== FastAPI App ====================

app = FastAPI(title="Resume Service", version="1.0.0")
resume_service = ResumeService()

allowed_origins = [origin.strip() for origin in FRONTEND_ORIGINS.split(",") if origin.strip()]
app.add_middleware(
    CORSMiddleware,
    allow_origins=allowed_origins or ["http://localhost:3000"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)

async def get_current_user_id() -> UUID:
    """In production, extract from JWT token"""
    return uuid4()

@app.post("/v1/resumes", status_code=status.HTTP_201_CREATED)
async def create_resume(
    resume_data: ResumeCreate,
    db: AsyncSession = Depends(get_db),
    background_tasks: BackgroundTasks = None,
    user_id: UUID = Depends(get_current_user_id)
):
    resume = Resume(
        user_id=user_id,
        source_type=resume_data.source_type,
        source_url=resume_data.source_url,
        raw_content="",
        is_active=True
    )
    db.add(resume)
    await db.flush()
    return {"id": str(resume.id), "status": "created"}

@app.post("/v1/resumes/{resume_id}/parse")
async def parse_resume(
    resume_id: UUID,
    payload: ResumeParseRequest,
    db: AsyncSession = Depends(get_db),
    background_tasks: BackgroundTasks = None
):
    result = await db.execute(select(Resume).where(Resume.id == resume_id))
    resume = result.scalar_one_or_none()
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")

    parsed_data = await resume_service.parse_resume(resume_id, payload.raw_content)

    resume.raw_content = payload.raw_content
    resume.parsed_data = parsed_data

    for section_type, content in parsed_data.items():
        if isinstance(content, (dict, list)):
            section = ResumeSection(
                resume_id=resume.id,
                section_type=section_type,
                content=content,
                ai_extracted=True,
                confidence_score=0.95
            )
            db.add(section)

    await db.commit()
    return {"id": str(resume.id), "parsed_data": parsed_data, "status": "parsed"}

@app.get("/v1/resumes/{resume_id}")
async def get_resume(resume_id: UUID, db: AsyncSession = Depends(get_db)):
    result = await db.execute(select(Resume).where(Resume.id == resume_id))
    resume = result.scalar_one_or_none()
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")

    sections_result = await db.execute(
        select(ResumeSection).where(ResumeSection.resume_id == resume_id)
    )
    sections = sections_result.scalars().all()

    return {
        "id": str(resume.id),
        "user_id": str(resume.user_id),
        "source_type": resume.source_type.value,
        "parsed_data": resume.parsed_data,
        "ats_score": resume.ats_score,
        "improvement_suggestions": resume.improvement_suggestions,
        "version": resume.version,
        "sections": [
            {"id": str(s.id), "type": s.section_type, "content": s.content, "confidence": s.confidence_score}
            for s in sections
        ],
        "created_at": resume.created_at.isoformat()
    }

@app.post("/v1/resumes/{resume_id}/analyze")
async def analyze_resume(
    resume_id: UUID,
    payload: ResumeAnalyzePayload,
    db: AsyncSession = Depends(get_db)
):
    result = await db.execute(select(Resume).where(Resume.id == resume_id))
    resume = result.scalar_one_or_none()
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    if not resume.parsed_data:
        raise HTTPException(status_code=400, detail="Resume not parsed yet")

    analysis = await resume_service.calculate_ats_score(resume.parsed_data, payload.target_role)
    resume.ats_score = analysis["overall"]
    resume.improvement_suggestions = analysis["improvements"]
    await db.commit()
    return {"resume_id": str(resume_id), "analysis": analysis}

@app.get("/v1/resumes/{resume_id}/improvements")
async def get_improvements(resume_id: UUID, db: AsyncSession = Depends(get_db)):
    result = await db.execute(select(Resume).where(Resume.id == resume_id))
    resume = result.scalar_one_or_none()
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    return {
        "resume_id": str(resume_id),
        "improvements": resume.improvement_suggestions or [],
        "ats_score": resume.ats_score
    }

@app.put("/v1/resumes/{resume_id}/sections/{section_type}")
async def update_section(
    resume_id: UUID,
    section_type: str,
    section_data: Dict[str, Any],
    db: AsyncSession = Depends(get_db)
):
    result = await db.execute(
        select(ResumeSection).where(
            ResumeSection.resume_id == resume_id,
            ResumeSection.section_type == section_type
        )
    )
    section = result.scalar_one_or_none()
    if section:
        section.content = section_data
        section.ai_extracted = False
    else:
        section = ResumeSection(
            resume_id=resume_id,
            section_type=section_type,
            content=section_data,
            ai_extracted=False
        )
        db.add(section)
    await db.commit()
    return {"status": "updated", "section_type": section_type}

@app.delete("/v1/resumes/{resume_id}")
async def delete_resume(resume_id: UUID, db: AsyncSession = Depends(get_db)):
    result = await db.execute(select(Resume).where(Resume.id == resume_id))
    resume = result.scalar_one_or_none()
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    resume.is_active = False
    await db.commit()
    return {"status": "deleted"}

@app.post("/api/upload-resume")
async def upload_resume(
    resume: UploadFile = File(...),
    user_id: Optional[str] = Form(None),
    db: AsyncSession = Depends(get_db)
):
    """
    Upload and analyze a resume file with STRICT validation.
    """
    try:
        # Step 1: File type validation
        if not resume_service.validate_file_type(resume.filename):
            raise HTTPException(
                status_code=400,
                detail="Invalid file type. Please upload a PDF, DOC, or DOCX file."
            )

        # Step 2: Read file content
        content = await resume.read()

        if len(content) < 100:
            raise HTTPException(status_code=400, detail="File is too small. Please upload a valid resume file.")

        if len(content) > 5 * 1024 * 1024:  # 5MB
            raise HTTPException(status_code=400, detail="File is too large. Maximum size is 5MB.")

        # Step 3: Extract text
        text = resume_service.extract_text_from_file(content, resume.filename)
        preview = text[:500].replace("\n", " ")
        print(f"[resume-upload] file={resume.filename} chars={len(text)} preview={preview}")

        if not text or len(text.strip()) < MIN_TEXT_LENGTH:
            raise HTTPException(
                status_code=400,
                detail="Could not extract readable text from file. Ensure it's a text-based PDF/DOCX, not a scanned image or encrypted file."
            )

        # Step 4: Content validation
        text_lower = text.lower()
        spam_indicators = [
            "verification code", "otp", "click here to verify",
            "security code", "confirm your email",
        ]
        if any(indicator in text_lower for indicator in spam_indicators):
            raise HTTPException(
                status_code=400,
                detail="This appears to be a verification email or security document, not a professional resume."
            )

        has_email = bool(re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text))
        has_phone = bool(re.search(r'[\+\(]?[1-9][0-9 .\-\(\)]{8,}[0-9]', text))

        section_indicators = {
            'experience': ['experience', 'employment', 'work history', 'professional experience', 'career'],
            'education': ['education', 'degree', 'university', 'college', 'bachelor', 'master', 'phd'],
            'skills': ['skills', 'technologies', 'proficiencies', 'technical skills', 'competencies']
        }

        matched_core_sections = sum(
            1 for keywords in section_indicators.values()
            if any(keyword in text_lower for keyword in keywords)
        )

        if not (has_email and has_phone and matched_core_sections >= 2):
            raise HTTPException(
                status_code=400,
                detail="Document does not appear to be a professional resume. Missing required sections: contact info, experience, education, or skills."
            )

        non_alpha_non_space_chars = sum(1 for c in text if not c.isalpha() and not c.isspace())
        noise_ratio = non_alpha_non_space_chars / max(len(text), 1)
        if noise_ratio > 0.40:
            raise HTTPException(
                status_code=400,
                detail="Could not extract readable text from file. The document appears too noisy or non-textual."
            )

        # Step 5: Parse resume
        resume_id = uuid4()
        parsed = await resume_service.parse_resume(resume_id, text)
        print(
            f"[resume-parse] parsed_keys={list(parsed.keys())} "
            f"skills_count={len(resume_service._extract_skill_names(parsed))}"
        )

        # Step 6: Generate ATS analysis
        analysis_result = await resume_service.generate_professional_ats_analysis(parsed)

        # Step 7: Store in database (non-fatal on failure)
        try:
            user_uuid = uuid4() if not user_id else UUID(user_id)
            resume_analysis = ResumeAnalysis(
                user_id=user_uuid,
                file_name=resume.filename,
                ats_score=analysis_result["atsScore"],
                skills_detected=analysis_result["skillsDetected"],
                missing_skills=analysis_result["missingSkills"],
                strengths=analysis_result["strengths"],
                improvements=analysis_result["improvements"],
                suggestions=analysis_result["suggestions"],
                uploaded_at=datetime.utcnow()
            )
            db.add(resume_analysis)
            await db.commit()
        except Exception as db_error:
            print(f"Warning: Could not save to database: {db_error}")

        # Step 8: Return response
        return {
            "ats_score": analysis_result["atsScore"],
            "skills_detected": analysis_result["skillsDetected"],
            "missing_skills": analysis_result["missingSkills"],
            "strengths": analysis_result["strengths"],
            "improvements": analysis_result["improvements"],
            "suggestions": analysis_result["suggestions"],
            "professional_feedback": analysis_result.get(
                "professionalFeedback",
                _generate_analysis_summary(analysis_result["atsScore"])
            ),
            "analysis": analysis_result.get(
                "professionalFeedback",
                _generate_analysis_summary(analysis_result["atsScore"])
            )
        }

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing resume: {str(e)}")


def _generate_analysis_summary(ats_score: float) -> str:
    if ats_score >= 80:
        return "Your resume is excellent! It has strong keywords, good formatting, and relevant experience. You're well-positioned for ATS screening."
    elif ats_score >= 60:
        return "Your resume is good but could use some improvements. Consider adding more relevant keywords and quantifiable achievements."
    else:
        return "Your resume needs work. Focus on adding more technical keywords, improving formatting, and including measurable achievements."


@app.get("/api/resume-history")
async def get_resume_history(
    user_id: Optional[str] = None,
    db: AsyncSession = Depends(get_db)
):
    try:
        user_uuid = UUID(user_id) if user_id else None

        query = select(ResumeAnalysis).order_by(desc(ResumeAnalysis.uploaded_at)).limit(20)
        if user_uuid:
            query = query.where(ResumeAnalysis.user_id == user_uuid)

        result = await db.execute(query)
        analyses = result.scalars().all()

        return {
            "history": [
                {
                    "id": str(a.id),
                    "file_name": a.file_name,
                    "ats_score": a.ats_score,
                    "skills_detected": a.skills_detected or [],
                    "missing_skills": a.missing_skills or [],
                    "strengths": a.strengths or [],
                    "improvements": a.improvements or [],
                    "suggestions": a.suggestions or [],
                    "uploaded_at": a.uploaded_at.isoformat() if a.uploaded_at else None
                }
                for a in analyses
            ],
            "total": len(analyses)
        }
    except Exception as e:
        return {"history": [], "total": 0, "error": str(e)}


@app.get("/health")
async def health_check():
    return {
        "status": "healthy",
        "service": "resume",
        "ai_configured": bool(OPENROUTER_API_KEY),
    }


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8002)